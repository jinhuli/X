from docx import Document


def writeTable(data,document):
    """将dataframe写入word中"""
    a = len(data)
    b = len(data.columns)
    table = document.add_table(rows=1, cols=b + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "index"

    for cols in range(b):
        hdr_cells[cols + 1].text = data.columns[cols]

    for i in range(a):
        row_cells = table.add_row().cells
        row_cells[0].text = str(data.index[i])
        print("写入第%s行数据"%i)

        for j in range(b):
            if type(data.iloc[i, j]) == pd.Timestamp:
                row_cells[j + 1].text = data.iloc[i, j].strftime('%Y-%m-%d')
            elif type(data.iloc[i, j]) == str:
                row_cells[j + 1].text = data.iloc[i, j]
            elif type(data.iloc[i, j]) == float:
                row_cells[j + 1].text = str(np.round(data.iloc[i, j],2))

            else:
                print("数据包含异常格式变量")



document = Document()  #创建文档类
document.add_heading("中证500指数分析", 0)
p = document.add_paragraph('本报告由python程序自动生成，数据来自于WIND')
document.add_heading('净值走势')
from docx.shared import Inches
document.add_picture("E:\\github\\X\\2.jpg", width=Inches(4.0))

document.add_heading('因子相关性分析')

document.add_heading('指数收益归因')

document.add_heading('行业归因', level=2)

document.add_heading('风格归因', level=2)

document.save("E:\\github\\X\\demo.docx")