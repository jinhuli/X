from WindPy import *
from datetime import *
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib as mpl
# 解决中文乱码问题
#sans-serif就是无衬线字体，是一种通用字体族。
#常见的无衬线字体有 Trebuchet MS, Tahoma, Verdana, Arial, Helvetica, 中文的幼圆、隶书等等。
mpl.rcParams['font.sans-serif']=['SimHei'] #指定默认字体 SimHei为黑体
mpl.rcParams['axes.unicode_minus']=False #用来正常显示负号
plt.style.use('ggplot')

from docx import Document
import matplotlib
matplotlib.matplotlib_fname()
w.start()


def writeTable(data,document):

    a = len(data)
    b = len(data.columns)
    table = document.add_table(rows=1, cols=b + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "index"

    for cols in range(b):
        hdr_cells[cols + 1].text = data.columns[cols]

    for i in range(a):
        row_cells = table.add_row().cells
        row_cells[0].text = str(data.index[i])
        print("写入第%s行数据"%i)
        for j in range(b):
            if type(data.iloc[i, j]) == pd.Timestamp:
                row_cells[j + 1].text = data.iloc[i, j].strftime('%Y-%m-%d')
            elif type(data.iloc[i, j]) == str:
                row_cells[j + 1].text = data.iloc[i, j]
            elif type(data.iloc[i, j]) == float:
                row_cells[j + 1].text = str(np.round(data.iloc[i, j],2))

            else:
                print("数据包含异常格式变量")

###指数截面表现#####
class index(object):
    def __init__(self, tradate,indexCode ,benchmark):
        self.indexCode = indexCode
        self.tradate = tradate
        self.benchmark = benchmark
        self.factor = None
        self.data = None
        self.indextimeseries = None
        codelist = w.wset('indexconstituent', "date=" + self.tradate + ";windcode=" + self.indexCode)
        self.list = pd.DataFrame(codelist.Data, columns=codelist.Codes, index=codelist.Fields,dtype=float).T["wind_code"].tolist()

    def get_code_list(self):
        """获取指数成分"""
        codelist = w.wset("indexconstituent","date="+self.tradate+";windcode="+self.indexCode)
        self.codelist = pd.DataFrame(codelist.Data, columns=codelist.Codes, index=codelist.Fields,dtype=float).T
        return self.codelist

    def get_factor(self):
        """获取指数截面因子"""
        tradedate = datetime.strptime(self.tradate, '%Y-%m-%d').strftime('%Y%m%d')
        startdate = w.tdaysoffset(-1, "2018-10-10", "Period=M").Data[0][0].strftime('%Y%m%d')
        year = datetime.strptime(self.tradate, '%Y-%m-%d').year
        ind='fa_roicebit_ttm,fa_ocftoor_ttm,fa_debttoasset,fa_npgr_ttm,fa_orgr_ttm,tech_price1y,pe_ttm,val_mvtoebitda_ttm,pb_lf,beta_24m,annualstdevr_24m'
        factor =w.wss(self.list, ind, "tradeDate=%s"%(tradedate))
        self.factor = pd.DataFrame(factor.Data, columns=factor.Codes, index=factor.Fields,dtype=float).T
        return self.factor

    def get_data(self):
        codelist = self.get_code_list()
        codelist.set_index(['wind_code'], inplace = True)

        factor = self.get_factor()
        self.data = pd.concat([codelist,factor], axis=1, join_axes=[factor.index])
        return self.data

    def get_indextimeseries(self):
        data = w.wsd(self.indexCode,"close,pct_chg", "ED-5Y",self.tradate, "")
        self.indextimeseries = pd.DataFrame(data.Data, columns=data.Times, index=data.Fields,dtype=float).T
        self.indextimeseries["PCT_CHG_1"] = self.indextimeseries["PCT_CHG"].map(lambda x: x / 100 + 1)
        self.indextimeseries["nav"] = self.indextimeseries["PCT_CHG_1"].cumprod()





##因子分组函数
def cla(n, lim):
    return '[%.f ,%.f)' % (lim * ((n-0.01 )// lim), lim * ((n-0.01 ) // lim) + lim)

  ##  b = data["VAL_FLOATMV"].apply(cla, args=(100,)).values
  ##  data["%s_3" % factor.upper()] = b

def visual(data,title, x=None, y=None):
    fig, ax = plt.subplots()
    ax.plot(x, y)
    plt.show()




if __name__ == '__main__':
    ####准备数据
    zz500=index("2018-9-26", "000905.SH", "000905.SH")
    zz500.get_code_list()
    zz500.get_factor()
    zz500.get_data()
    zz500.get_indextimeseries()
    timeseries = zz500.indextimeseries
    data = zz500.get_data()
    ##净值走势图
    aa=plt.figure()
    timeseries["nav"].plot()
    plt.savefig("E:\\github\\X\\2.jpg")

    document = Document()  #创建文档类
    document.add_heading("中证500指数分析", 0)
    p = document.add_paragraph('本报告由python程序自动生成，数据来自于WIND')
    document.add_heading('净值走势')
    from docx.shared import Inches
    document.add_picture("E:\\github\\X\\2.jpg", width=Inches(4.0))

    #####常用因子######

    ##价值类##
    ##fa_roicebit_ttm                 投入资本回报率ROIC
    ##fa_ocftoor_ttm                  经营活动产生的现金流量净额/营业收入
    ##fa_debttoasset                  资产负债率
    ##fa_npgr_ttm                     净利润增长率
    ##fa_orgr_ttm                     营业收入增长率
    ##pe_ttm                          市盈率
    ##val_mvtoebitda_ttm              市值/EBITDA
    ##pb_lf                           市净率


    ##风险类##
    ##beta_24m                        BETA近24个月
    ##annualstdevr_24m                年化波动率近24个月

    ##量价类##
    ##tech_price1y                   当前股价/过去一年均价-1

    ###因子权重分析

    factor_i = ['PCT_CHG_PER', 'VAL_PE_DEDUCTED_TTM', 'PB_LYR', 'DIVIDENDYIELD2', 'PE_EST', 'EST_PEG',
                'BOLL', 'DMA', 'VAL_LNFLOATMV', 'VAL_FLOATMV']
    data[factor_i] = data[factor_i].astype(float)
    factor_d = []
    document.add_heading('因子权重分析')
      ###因子分组均分成5组
    for i in factor_i:
        data[i+"_group"] = data[i].rank(ascending=False).apply(cla, args=(100,)).values

    from pandas.plotting import scatter_matrix
    scatter_matrix(data[['PCT_CHG_PER', 'PE_EST', 'EST_PEG', 'BOLL', 'DMA', 'VAL_LNFLOATMV']], alpha=0.2, figsize=(6, 6), diagonal='kde')

    pie_data1 = data.groupby("INDUSTRY_SW")["i_weight"].sum().sort_values()
    fig1, ax1 = plt.subplots()
    ax1.pie(pie_data1,labels=pie_data1.index.tolist(),shadow=True, autopct='%1.1f%%', startangle=90)
    ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
    plt.title("因子权重")
    plt.show()
    plt.savefig("E:\\github\\X\\1.jpg")
    document.add_picture("E:\\github\\X\\1.jpg", width=Inches(4.0))

    fig, axs = plt.subplots(5, 2, figsize=(5, 5))
    for i in range(len(factor_i)):
        factor=factor_i[i]
        pie_data = data.groupby(factor + "_group")["i_weight"].sum().sort_values()
        a=i//2
        b=i%2
        axs[a, b].pie(pie_data, labels=pie_data.index.tolist(), shadow=True, autopct='%1.1f%%', startangle=90)
        axs[a, b].set_title(factor)
    plt.show()









    ###因子交叉分析
    dddata = data.groupby( ['VAL_PE_DEDUCTED_TTM_group','EST_PEG_group'])["i_weight"].sum()
    x= dddata.index.levels[1].tolist()
    y = dddata.index.levels[0].tolist()
    X,Y =np.meshgrid(x,y)
    def zz(x,y):
        return dddata[x][y]

    Z=dddata.map(zz)





    document.add_picture("E:\\github\\X\\3.jpg", width=Inches(4.0))

    ##分组

    ##writeTable(data.head(),document)


    ##因子相关性
    document.add_heading('因子相关性分析')

    #####收益归因--分组收益统计####
    document.add_heading('指数收益归因')

    document.add_heading('行业归因', level = 2)

    document.add_heading('风格归因', level = 2)

    document.save("E:\\github\\X\\demo.docx")




